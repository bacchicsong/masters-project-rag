import json
import os
import PyPDF2
from docx import Document
from typing import List, Dict, Any

class FileParser:
    def __init__(self):
        self.supported_files = [
            "Antti_Ilmanen_Expected_Returns_A.pdf",
            "Hoi,_Steven_C_H___Li,_Bin_Online.pdf",
            "Hull_J_C_-Options_Futures_and_Other_Derivatives_9th_edition.pdf",
            "MOEX_bonds.pdf",
            "MOEX_futures_options.pdf",
            "MOEX_indexes.docx",
            "Richard Grinold Active_Portfolio_Management_A_quantitative_approach_for_providing.pdf"
        ]
    
    def parse_pdf(self, file_path: str) -> Dict[str, Any]:
        """Парсинг PDF файлов"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Извлекаем метаданные
                metadata = pdf_reader.metadata
                metadata_dict = {
                    'title': getattr(metadata, 'title', ''),
                    'author': getattr(metadata, 'author', ''),
                    'subject': getattr(metadata, 'subject', ''),
                    'creator': getattr(metadata, 'creator', ''),
                    'producer': getattr(metadata, 'producer', ''),
                    'creation_date': str(getattr(metadata, 'creation_date', '')),
                    'modification_date': str(getattr(metadata, 'modification_date', ''))
                }
                
                # Извлекаем текст со всех страниц
                text_content = []
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    text_content.append({
                        'page': page_num + 1,
                        'text': text
                    })
                
                return {
                    'file_name': os.path.basename(file_path),
                    'file_type': 'PDF',
                    'metadata': metadata_dict,
                    'page_count': len(pdf_reader.pages),
                    'content': text_content
                }
                
        except Exception as e:
            return {
                'file_name': os.path.basename(file_path),
                'file_type': 'PDF',
                'error': f"Ошибка при чтении PDF: {str(e)}"
            }
    
    def parse_docx(self, file_path: str) -> Dict[str, Any]:
        """Парсинг DOCX файлов"""
        try:
            doc = Document(file_path)
            
            # Извлекаем метаданные
            core_props = doc.core_properties
            metadata_dict = {
                'title': getattr(core_props, 'title', ''),
                'author': getattr(core_props, 'author', ''),
                'subject': getattr(core_props, 'subject', ''),
                'keywords': getattr(core_props, 'keywords', ''),
                'comments': getattr(core_props, 'comments', ''),
                'created': str(getattr(core_props, 'created', '')),
                'modified': str(getattr(core_props, 'modified', ''))
            }
            
            # Извлекаем текст из параграфов
            paragraphs = []
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():  # Игнорируем пустые параграфы
                    paragraphs.append({
                        'paragraph_number': i + 1,
                        'text': paragraph.text
                    })
            
            # Извлекаем текст из таблиц
            tables_content = []
            for table_num, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables_content.append({
                    'table_number': table_num + 1,
                    'data': table_data
                })
            
            return {
                'file_name': os.path.basename(file_path),
                'file_type': 'DOCX',
                'metadata': metadata_dict,
                'paragraph_count': len(paragraphs),
                'table_count': len(tables_content),
                'paragraphs': paragraphs,
                'tables': tables_content
            }
            
        except Exception as e:
            return {
                'file_name': os.path.basename(file_path),
                'file_type': 'DOCX',
                'error': f"Ошибка при чтении DOCX: {str(e)}"
            }
    
    def parse_all_files(self) -> List[Dict[str, Any]]:
        """Парсинг всех поддерживаемых файлов"""
        results = []
        
        for file_name in self.supported_files:
            if not os.path.exists(file_name):
                print(f"Файл {file_name} не найден в текущей директории")
                results.append({
                    'file_name': file_name,
                    'error': 'Файл не найден'
                })
                continue
            
            print(f"Обработка файла: {file_name}")
            
            if file_name.endswith('.pdf'):
                result = self.parse_pdf(file_name)
            elif file_name.endswith('.docx'):
                result = self.parse_docx(file_name)
            else:
                result = {
                    'file_name': file_name,
                    'error': 'Неподдерживаемый формат файла'
                }
            
            results.append(result)
        
        return results
    
    def save_to_json(self, data: List[Dict[str, Any]], output_file: str = 'parsed_files.json'):
        """Сохранение результатов в JSON файл"""
        try:
            with open(output_file, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
            print(f"Результаты сохранены в файл: {output_file}")
        except Exception as e:
            print(f"Ошибка при сохранении JSON: {str(e)}")

def main():
    # Проверяем наличие необходимых библиотек
    try:
        import PyPDF2
        from docx import Document
    except ImportError as e:
        print("Необходимо установить требуемые библиотеки:")
        print("pip install PyPDF2 python-docx")
        return
    
    # Создаем парсер и обрабатываем файлы
    parser = FileParser()
    
    print("Начало обработки файлов...")
    parsed_data = parser.parse_all_files()
    
    # Сохраняем результаты
    parser.save_to_json(parsed_data)
    
    # Выводим статистику
    successful = len([f for f in parsed_data if 'error' not in f])
    print(f"\nОбработка завершена:")
    print(f"Успешно обработано: {successful}/{len(parser.supported_files)} файлов")

if __name__ == "__main__":
    main()
